import os
import time
import logging
import re
import tempfile # Import tempfile for secure temporary file handling
import shutil # Import shutil for file copying
from io import BytesIO # Import BytesIO for in-memory file handling

from dotenv import load_dotenv
from flask import Flask, request, jsonify, render_template, session, redirect, url_for, send_from_directory
import pdfplumber
from docx import Document # Already imported
from docx.shared import Inches # For potential layout control in DOCX
import nltk
from nltk.corpus import stopwords
import stripe
from openai import OpenAI

# Load environment variables from .env file
load_dotenv()

# Logging setup
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Download stopwords once (only if not already downloaded)
try:
    nltk.data.find('corpora/stopwords')
except nltk.downloader.DownloadError:
    logger.info("NLTK stopwords not found, downloading...")
    nltk.download('stopwords', quiet=True)
    logger.info("NLTK stopwords downloaded.")

# Flask app
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'a_very_secret_key_that_should_be_changed_in_production')

# Ensure a static folder exists for serving files
if not os.path.exists('static'):
    os.makedirs('static')
    logger.info("Created 'static' directory.")

# Stripe API keys (ensure these are in your .env file)
stripe.api_key = os.getenv('STRIPE_SECRET_KEY')
STRIPE_PUB_KEY = os.getenv('STRIPE_PUBLISHABLE_KEY')

# OpenAI client (ensure OPENAI_API_KEY is in your .env file)
try:
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    logger.info("OpenAI client initialized.")
except Exception as e:
    logger.error(f"Failed to initialize OpenAI client: {e}. Please check OPENAI_API_KEY in your .env file.")
    client = None # Set client to None if initialization fails

# Text extraction utilities
def extract_text_from_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        logger.error(f"Error extracting text from PDF {path}: {e}")
        raise ValueError("Could not extract text from PDF. Is it a valid PDF?")

def extract_text_from_docx(path):
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {path}: {e}")
        raise ValueError("Could not extract text from DOCX. Is it a valid DOCX?")

# Simple regex-based tokenizer + stopwords
def preprocess(text):
    tokens = re.findall(r"\b[a-zA-Z]+\b", text.lower())
    stops = set(stopwords.words('english'))
    return [t for t in tokens if t not in stops]

# ATS scoring
def calculate_ats_score(resume_text, job_desc, premium=False):
    r_tokens = set(preprocess(resume_text))
    j_tokens = set(preprocess(job_desc))
    score = (len(r_tokens & j_tokens) / len(j_tokens) * 100) if j_tokens else 0
    score = min(round(score, 2), 100)
    suggestions = []
    if not premium:
        suggestions = ["Upgrade to Premium for detailed suggestions."]
    else:
        missing = list(j_tokens - r_tokens)
        if missing:
            suggestions.append(f"Consider adding keywords from the job description: {', '.join(missing[:5])}.")
        if score < 80:
            suggestions.append("Quantify your achievements with numbers and metrics (e.g., 'Increased sales by 20%').")
        if len(resume_text.split()) > 700: # Adjusted word count for a typical 1-2 page resume
            suggestions.append("Concise resumes are often preferred. Try to shorten your resume to focus on key accomplishments.")
        if not any(re.search(r'\b(managed|led|developed|implemented|achieved)\b', resume_text, re.IGNORECASE) for _ in range(1)):
             suggestions.append("Use strong action verbs at the beginning of your bullet points.")
    return score, suggestions

# Summary generation
def generate_summary(resume_text, job_desc, premium=False):
    if not premium:
        return "Upgrade to Premium for a tailored summary."
    
    # Use OpenAI to generate a concise summary
    if client is None:
        return "AI summary unavailable due to OpenAI client initialization error."

    prompt = f"""
    Based on the following resume text and job description, generate a concise, impactful, one-sentence professional summary for the resume.
    Focus on aligning the summary with the key requirements and keywords from the job description.

    Job Description:
    \"\"\"{job_desc}\"\"\"

    Resume Text (first 1000 chars):
    \"\"\"{resume_text[:1000]}...\"\"\"

    Generated Summary:
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o", # Using gpt-4o for better quality
            messages=[
                {"role": "system", "content": "You are an expert resume writer, crafting concise and impactful professional summaries."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=60
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"OpenAI summary generation error: {e}")
        return "AI summary generation failed. Please try again later."

# Detailed AI feedback via OpenAI
def generate_detailed_feedback(resume_text: str, job_desc: str, score: float) -> str:
    if client is None:
        return "Detailed AI feedback unavailable due to OpenAI client initialization error."

    prompt = f"""
You are an expert resume coach. A user's resume received an ATS match score of {score:.2f}% against this job description:

Job Description:
\"\"\"{job_desc}\"\"\"

Resume Text (first 2000 chars):
\"\"\"{resume_text[:2000]}...\"\"\"

Please provide:
1. A clear, concise diagnosis of why the score is what it is.
2. 3–5 concrete, prioritized improvements to make the resume better match the job (e.g., missing keywords, rewording, formatting, quantifiable accomplishments).
3. A one-sentence highlight the user can add to their summary.

Be direct, helpful, and practical. Focus on match relevance and keyword alignment.
Format the improvements as a numbered list.
"""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o", # Using gpt-4o for better quality
            messages=[
                {"role":"system","content":"You help users improve resumes for ATS."},
                {"role":"user","content":prompt}
            ],
            temperature=0.7,
            max_tokens=500
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"OpenAI detailed feedback error: {e}")
        return "Detailed feedback is unavailable right now."

def generate_tailored_resume(original_resume_text: str, job_desc: str, output_dir: str) -> str:
    """
    Generates a tailored resume text using OpenAI and saves it as a new DOCX file.
    This function will generate a new document with tailored text, but will NOT retain
    the original resume's complex formatting (e.g., columns, specific fonts, etc.).
    It will be a simple DOCX file.
    """
    logger.info("Generating tailored resume text via OpenAI...")
    if client is None:
        logger.error("OpenAI client not initialized, cannot generate tailored resume text.")
        return None

    prompt = f"""
    You are an expert resume writer. Your task is to tailor a resume to a specific job description.
    Rewrite the provided resume text to maximize its relevance to the job description.
    
    IMPORTANT GUIDELINES:
    1.  **DO NOT invent new projects, experiences, or qualifications.** Only rephrase or emphasize existing information from the original resume.
    2.  Focus on incorporating keywords from the job description and highlighting relevant achievements already present in the original resume.
    3.  **Keep the tailored resume concise, aiming for a single page.** Prioritize the most relevant information and condense where possible.
    4.  If you cannot incorporate a specific keyword without inventing content, do not include it.
    5.  The output should be the full tailored resume content, ready to be put into a document.
    
    Job Description:
    \"\"\"{job_desc}\"\"\"

    Original Resume Text:
    \"\"\"{original_resume_text}\"\"\"

    Tailored Resume (start directly with the resume content, no introductory phrases):
    """
    try:
        resp = client.chat.completions.create(
            model="gpt-4o", # Using gpt-4o for better quality
            messages=[
                {"role": "system", "content": "You are an expert resume writer, tailoring resumes to job descriptions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500 # Allow enough tokens for a one-page resume
        )
        tailored_text = resp.choices[0].message.content.strip()

        # Create a new DOCX document
        document = Document()
        
        # Add paragraphs from the tailored text, splitting by newlines for basic structure
        for paragraph_text in tailored_text.split('\n'):
            if paragraph_text.strip(): # Only add non-empty paragraphs
                document.add_paragraph(paragraph_text.strip())

        # Save the document to a temporary in-memory file
        byte_io = BytesIO()
        document.save(byte_io)
        byte_io.seek(0) # Rewind to the beginning of the stream

        # Define a filename for the tailored resume
        tailored_filename = f"tailored_resume_{int(time.time())}.docx"
        tailored_filepath = os.path.join(output_dir, tailored_filename)

        # Save the in-memory DOCX to the static folder
        with open(tailored_filepath, 'wb') as f:
            f.write(byte_io.getvalue())

        logger.info(f"Tailored resume generated and saved to {tailored_filepath}")
        return tailored_filename
    except Exception as e:
        logger.error(f"Error generating or saving tailored resume: {e}", exc_info=True)
        return None


# Routes
@app.route('/')
def index():
    return render_template('index.html', stripe_pub_key=STRIPE_PUB_KEY)

@app.route('/scan', methods=['POST'])
def scan():
    logger.info("Received scan request.")
    temp_resume_path = None # Initialize to None
    try:
        if 'resume' not in request.files or not request.form.get('job_description'):
            logger.warning("Missing resume file or job description.")
            return jsonify(error='Missing resume file or job description'), 400

        # Make all features free and unlimited by forcing premium to True
        premium = True # Forcing premium for testing

        f = request.files['resume']
        job_description = request.form['job_description']
        ext = f.filename.rsplit('.', 1)[-1].lower()

        if ext not in ('pdf', 'docx'):
            logger.warning(f"Unsupported file type: {ext}")
            return jsonify(error='Unsupported file type. Please upload a PDF or DOCX.'), 400

        # Use tempfile to create a secure temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as temp_resume_file:
            f.save(temp_resume_file.name)
            temp_resume_path = temp_resume_file.name
        
        logger.info(f"Resume saved temporarily to: {temp_resume_path}")

        try:
            resume_text = extract_text_from_pdf(temp_resume_path) if ext == 'pdf' else extract_text_from_docx(temp_resume_path)
            if not resume_text.strip():
                raise ValueError("Could not extract any meaningful text from the resume.")
        except ValueError as ve:
            logger.error(f"Text extraction failed: {ve}")
            return jsonify(error=str(ve)), 400
        except Exception as e:
            logger.error(f"Unexpected error during text extraction: {e}")
            return jsonify(error="Failed to process resume file. Please ensure it's not corrupted."), 400
        finally:
            # Ensure the temporary file is deleted if it was created
            if temp_resume_path and os.path.exists(temp_resume_path):
                os.remove(temp_resume_path)
                logger.info(f"Temporary input file deleted: {temp_resume_path}")


        score, suggestions = calculate_ats_score(resume_text, job_description, premium)
        summary = generate_summary(resume_text, job_description, premium)

        ai_feedback = "Detailed feedback is unavailable right now."
        if client: # Only try if OpenAI client was successfully initialized
            try:
                ai_feedback = generate_detailed_feedback(resume_text, job_description, score)
            except Exception as e:
                logger.error(f"OpenAI error during detailed feedback generation: {e}")
        else:
            logger.warning("OpenAI client not initialized, skipping detailed feedback generation.")

        # Sample checklist data - in real implementation, calculate these dynamically
        checklist = [
            {"criteria": "Resume length under 2 pages", "passed": len(resume_text.split()) < 1000},
            {"criteria": "Includes relevant keywords", "passed": score >= 70},
            {"criteria": "Has quantifiable achievements (e.g., numbers, metrics)", "passed": bool(re.search(r'\d+%|\$\d+|increased|decreased|reduced|managed [0-9]+|over [0-9]+', resume_text, re.IGNORECASE))},
            {"criteria": "Proper formatting and structure (e.g., consistent headings)", "passed": True},  # Placeholder
            {"criteria": "No spelling or grammar errors", "passed": True}    # Placeholder
        ]

        # Generate tailored resume (now generates a new DOCX)
        tailored_resume_filename = generate_tailored_resume(resume_text, job_description, 'static')
        tailored_resume_url = url_for('static', filename=tailored_resume_filename) if tailored_resume_filename else None

        logger.info("Scan successful, returning results.")
        return jsonify(
            ats_score=score,
            suggestions=suggestions,
            summary=summary,
            ai_feedback=ai_feedback,
            checklist=checklist,
            tailored_resume_url=tailored_resume_url,
            is_free_user=False # Always False if premium is forced to True
        )
    except Exception as e:
        logger.error(f"Scan endpoint error: {e}", exc_info=True)
        return jsonify(error=f"An internal error occurred: {str(e)}. Please check server logs."), 500

@app.route('/create-checkout-session', methods=['POST'])
def checkout():
    logger.info("Received checkout session request.")
    try:
        sess = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency':'usd',
                    'product_data':{'name':'Premium Subscription'},
                    'unit_amount':999 # $9.99
                },
                'quantity':1
            }],
            mode='subscription',
            success_url=request.host_url + 'success',
            cancel_url=request.host_url + 'cancel'
        )
        return jsonify(id=sess.id)
    except Exception as e:
        logger.error(f"Stripe checkout error: {e}", exc_info=True)
        return jsonify(error=str(e)), 500

@app.route('/success')
def success():
    session['is_premium'] = True
    logger.info("Payment successful, user is now premium.")
    return redirect(url_for('index'))

@app.route('/cancel')
def cancel():
    logger.info("Payment cancelled.")
    return redirect(url_for('index'))

if __name__ == '__main__':
    logger.info("Starting Flask app...")
    # The app is typically run via run.py, but this block ensures it can be run directly if needed.
    app.run(host='127.0.0.1', port=5000, debug=True)
